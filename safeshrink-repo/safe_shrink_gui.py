"""
safe_shrink_gui.py - GUI 专用接口
封装 DocSlimmer 和 DocSanitizer，提供给 PyQt6 界面使用
"""

from pathlib import Path
from safe_shrink import DocSlimmer, DocSanitizer, read_file, write_file, DEPS

# SSD 转换为可选功能（在 frozen 环境中可能因 markitdown 依赖问题无法加载）
try:
    from format_to_ssd import convert_to_ssd_v2
except ImportError as e:
    import sys
    print(f"[WARNING] format_to_ssd 导入失败（禁用 SSD 转换功能）: {e}")
    convert_to_ssd_v2 = None

# ========== 注入新增脱敏类型到 DocSanitizer ==========
_EXTRA_PATTERNS = {
    '护照号': r'\b[EeGgPpDdSsHhLl][A-Za-z]?\d{7,9}\b',
    'Mac地址': r'\b([0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b',
    'IMEI': r'(?<!\d)\d{15}(?!\d)',
    '车牌号': r'[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤青藏川宁琼使领警][A-Z][·]?[A-HJ-NP-Z0-9]{4,5}[A-HJ-NP-Z0-9挂学港澳]?',
    '社保卡号': r'(?<!\d)\d{14,18}(?!\d)',
    '医保卡号': r'(?<!\d)\d{15,18}(?!\d)',
    '病历号': r'(?:(?:BL|MR|EMR|MZ|ZY|门诊号?|住院号?|病历号?|病案号?)[：:\-]?\d{4,}|\b[A-Z]{2,3}[-]?\d{6,12}\b)',
    '公文份号': r'(?:(?:No|NO|Nr|№)[-:\s]*\d{4,}|份号[：:\s]+\d+|第\d{2,4}[-]\d{4,}号|文件编号[：:\s]*[A-Z]{2,3}[-_]\d{2,4}[-_]\d{3,})',
    '公文密级': r'(?:【[绝密机密秘密内部]+】|绝密\s*[★☆]?\s*\d*\s*年?|[机密秘密内部]+(?:文件|资料|通知|信息))',
    '公文文号': r'[\u4e00-\u9fa5]+发〔\d{4}〕\s*(?:第\s*)?\d+(?:号)?|[\u4e00-\u9fa5]+政办发〔\d{4}〕\s*(?:第\s*)?\d+(?:号)?',
}
for _k, _v in _EXTRA_PATTERNS.items():
    DocSanitizer.PATTERNS[_k] = _v
    DocSanitizer.ITEM_LABELS[_k] = _k
# 覆盖旧手机号正则，支持分隔符格式
DocSanitizer.PATTERNS['手机号'] = r'(?<!\d)1[3-9]\d[\s\-]?\d{4}[\s\-]?\d{4}(?!\d)'
# 覆盖合同编号正则，支持 CONTRACT/XZ-HT 等格式
DocSanitizer.PATTERNS['合同编号'] = r'[A-Z]{2,8}[-_/](?:[A-Z0-9]+[-_/])*[A-Z0-9]*\d{4,}[A-Z0-9]*(?:[-_/][A-Z0-9]+)*'
# ===============================================

# 所有可选模块：缺失不影响核心功能
try:
    from sanitize_enhanced import sanitize_enhanced as _sanitize_enhanced, FAKE_AVAILABLE
    ENHANCED_SANITIZE = True
except ImportError:
    ENHANCED_SANITIZE = False
    FAKE_AVAILABLE = False

try:
    from sanitize_tags import sanitize_all_with_tags
    SANITIZE_TAGS = True
except ImportError:
    SANITIZE_TAGS = False
    sanitize_all_with_tags = None

try:
    from chinese_ner import sanitize as sanitize_chinese
    CHINESE_NER = True
except ImportError:
    CHINESE_NER = False
    sanitize_chinese = None


def slim_content(content, compression_rate=0.5, options=None):
    """
    文档减肥 - 压缩文本内容
    
    Args:
        content: 文本内容
        compression_rate: 压缩强度 (0.0 - 1.0)
        options: dict, 可选参数
            - remove_comments: 移除注释
            - remove_empty: 移除空行
            - shorten: 缩短长文本
            - remove_ai: 去除 AI 写作痕迹
    
    Returns:
        str: 压缩后的内容
    """
    opts = options or {}
    
    if not content or not content.strip():
        return content
    
    result = content
    
    # 基础清理
    if opts.get('remove_empty', True):
        # 移除连续空行，保留最多一个
        import re
        result = re.sub(r'\n{3,}', '\n\n', result)
        result = re.sub(r'\n\s*\n', '\n\n', result)
    
    if opts.get('remove_comments', False):
        # 移除注释（简单匹配）
        import re
        # Python 风格注释
        result = re.sub(r'#.*$', '', result, flags=re.MULTILINE)
        # C 风格注释
        result = re.sub(r'//.*$', '', result, flags=re.MULTILINE)
        result = re.sub(r'/\*.*?\*/', '', result, flags=re.DOTALL)
    
    # 使用 DocSlimmer 压缩
    slimmer = DocSlimmer()
    slim_result = slimmer.slim(
        result, 
        compression_rate=compression_rate,
        remove_ai=opts.get('remove_ai', False)
    )
    
    return slim_result['result']


def detect_sensitive(content, types=None, custom_patterns=None):
    """
    检测敏感信息（增强版，带校验和验证 + 上下文过滤）

    Args:
        content: 文本内容
        types: list, 要检测的类型列表，None 表示全部
        custom_patterns: dict, 自定义规则

    Returns:
        list: 检测到的敏感信息列表 [{type, match, position}, ...]
    """
    import re

    if not content:
        return []

    sanitizer = DocSanitizer()
    items = types if types else sanitizer.available_items()
    results = []
    seen = set()  # 去重

    for item_name in items:
        if item_name not in sanitizer.PATTERNS:
            continue
        pattern = sanitizer.PATTERNS[item_name]

        for m in re.finditer(pattern, content):
            val = m.group()
            keep = True

            # 身份证：校验和验证
            if item_name == '身份证':
                keep, _ = _validate_id_card(val)

            # 银行卡：Luhn 验证
            elif item_name == '银行卡':
                keep, _ = _validate_bank_card(val)

            if keep and val not in seen:
                results.append({
                    'type': item_name,
                    'match': val,
                    'position': (m.start(), m.end())
                })
                seen.add(val)

    # 邮编：上下文过滤
    if '邮编' in items:
        zip_pattern = r'(?<!\d)\d{6}(?!\d)'
        try:
            for m in re.finditer(zip_pattern, content):
                val = m.group()
                if val in seen:
                    continue
                # 上下文检查：前后20字符内有邮编相关词
                ctx_start = max(0, m.start() - 20)
                ctx_end = min(len(content), m.end() + 20)
                ctx = content[ctx_start:ctx_end]
                zip_kw = ['邮编', '邮政编码', 'ZIP', 'Postal', '地址', '寄至']
                if any(kw in ctx for kw in zip_kw):
                    results.append({
                        'type': '邮编',
                        'match': val,
                        'position': (m.start(), m.end())
                    })
                    seen.add(val)
        except:
            pass

    # 自定义敏感词检测
    if custom_patterns and custom_patterns.get('keywords'):
        keywords = [k.strip() for k in custom_patterns['keywords'].split(',') if k.strip()]
        for kw in keywords:
            for m in re.finditer(re.escape(kw), content):
                val = m.group()
                if val not in seen:
                    results.append({
                        'type': f'自定义:{kw}',
                        'match': val,
                        'position': (m.start(), m.end())
                    })
                    seen.add(val)

    return results


def _validate_id_card(id_str):
    """验证18位身份证校验和"""
    if not id_str or len(id_str) != 18:
        return False, "length"
    if not all(c.isdigit() or c in 'Xx' for c in id_str):
        return False, "format"
    weights = [7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2]
    check = '10X98765432'
    try:
        total = sum(int(id_str[i]) * weights[i] for i in range(17))
        if check[total % 11] != id_str[-1].upper():
            return False, "checksum"
        return True, ""
    except:
        return False, "calc"


def _validate_bank_card(card_str):
    """验证银行卡 Luhn 校验"""
    import re
    card = re.sub(r'\s', '', card_str)
    if not card.isdigit() or len(card) < 13 or len(card) > 19:
        return False, "format"
    digits = [int(d) for d in card]
    checksum = 0
    for i in range(len(digits) - 2, -1, -2):
        doubled = digits[i] * 2
        checksum += doubled if doubled < 10 else doubled - 9
    for i in range(len(digits) - 1, -1, -2):
        checksum += digits[i]
    return checksum % 10 == 0, ""



def sanitize_content(content, types=None, custom_patterns=None, mode='mask'):
    import sys
    print(f"[DEBUG] sanitize_content called, types={types!r}, mode={mode!r}, ENHANCED_SANITIZE={ENHANCED_SANITIZE}")
    """
    文档脱敏 - 去除敏感信息
    
    现在包含：
    - 身份证/银行卡校验和验证
    - 中文姓名识别（基于上下文）
    - 中文地址识别（基于上下文）
    - 假名化替换（Faker）
    """
    """
    文档脱敏 - 去除敏感信息

    Args:
        content: 文本内容
        types: list, 要脱敏的类型列表，None 表示全部
        custom_patterns: dict, 自定义规则 {"keywords": "词1,词2", "regex": "正则1"}
        mode: 'mask' (遮罩) | 'pseudo' (假名化)

    Returns:
        str: 脱敏后的内容
    """
    if not content:
        return content

    # 使用增强脱敏模块（带校验和验证 + 上下文过滤 + 假名化）
    if ENHANCED_SANITIZE:
        result = _sanitize_enhanced(
            text=content,
            items=types,
            mode=mode,
            use_validation=True,
            use_context=True
        )
        sanitized = result['result']
    else:
        sanitizer = DocSanitizer()
        sanitized = sanitizer.sanitize(content, items=types)['result']

    # 应用自定义脱敏规则
    if custom_patterns:
        if custom_patterns.get('keywords'):
            for kw in custom_patterns['keywords'].split(','):
                kw = kw.strip()
                if kw:
                    sanitized = sanitized.replace(kw, '***')

        if custom_patterns.get('regex'):
            import json
            try:
                patterns = json.loads(custom_patterns['regex'])
                if isinstance(patterns, list):
                    for p in patterns:
                        if isinstance(p, dict) and 'pattern' in p:
                            sanitized = re.sub(p['pattern'], p.get('replace', '***'), sanitized)
            except:
                for patt in custom_patterns['regex'].split(','):
                    patt = patt.strip()
                    if patt:
                        try:
                            sanitized = re.sub(patt, '***', sanitized)
                        except:
                            pass

    # 中文姓名和地址脱敏
    if ENHANCED_SANITIZE:
        sanitized, _ner_stats = sanitize_chinese(sanitized)

    # 如果使用语义标签模式，转换为标签
    if mode == 'tags':
        # 重新用语义标签脱敏
        sanitized, mapper = sanitize_all_with_tags(content, types=types)
        # 返回脱敏后的文本和映射信息
        return sanitized

    return sanitized




def process_file_gui(file_path, action, options=None):
    """
    处理单个文件（GUI 调用）
    根据文件类型自动选择最合适的处理方式
    
    Args:
        file_path: 文件路径
        action: 'slim' 或 'sanitize'
        options: dict, 处理选项
    
    Returns:
        dict: {success, content, error, stats}
    """
    opts = options or {}
    ext = Path(file_path).suffix.lower()
    
    try:
        # ===== 特殊格式：直接操作文件，不走文本流 =====
        
        # Word 深度清理
        if ext == '.docx' and action == 'slim' and opts.get('deep_clean', False):
            output_path = opts.get('output_path', file_path)
            result = clean_docx_deep(file_path, output_path, {
                'remove_empty_paragraphs': True,
                'remove_bullet_runs': True,
                'remove_non_image_shapes': True,
            })
            if result['success']:
                return {
                    'success': True,
                    'content': None,  # 直接写文件，不返回文本
                    'output_path': result['output_path'],
                    'stats': result['stats'],
                    'direct_write': True,
                }
            else:
                return {'success': False, 'error': result['error']}
        
        # PDF 元数据清理
        if ext == '.pdf' and action == 'slim':
            output_path = opts.get('output_path', file_path)
            result = clean_pdf_metadata(file_path, output_path)
            if result['success']:
                return {
                    'success': True,
                    'content': None,
                    'output_path': result.get('output_path'),
                    'skipped': result.get('skipped', False),
                    'stats': {},
                    'direct_write': True,
                }
            else:
                return {'success': False, 'error': result['error']}
        
        # ===== 通用格式：读取文本 → 处理 → 返回文本 =====
        
        content = read_file(file_path, opts)
        
        # 根据格式做预处理
        if action == 'slim':
            original_content = content
            
            # 格式专属清理
            if ext == '.md':
                content = clean_txt_md(content, opts, is_markdown=True)
            elif ext == '.txt':
                content = clean_txt_md(content, opts, is_markdown=False)
            elif ext == '.json':
                content = clean_json_content(content, opts)
            elif ext == '.csv':
                content = clean_csv_content(content, opts)
            elif ext in ('.html', '.htm'):
                content = clean_html_content(content, opts)
            
            # 使用 slim_content 进行压缩（传入 compression_rate）
            compression_rate = opts.get('compression_rate', 0.5)
            content = slim_content(content, compression_rate, opts)
            
            # 计算压缩率
            stats = {
                'original_length': len(original_content),
                'new_length': len(content),
                'compression': round((1 - len(content)/len(original_content)) * 100, 1) if original_content else 0
            }
            
            return {
                'success': True,
                'content': content,
                'stats': stats,
                'direct_write': False,
                'original_size': len(original_content.encode('utf-8')),
                'new_size': len(content.encode('utf-8')),
            }
            
        elif action == 'sanitize':
            types = opts.get('sanitize_items', None)
            result = sanitize_content(content, types)
            detected = detect_sensitive(content, types)
            stats = {
                'items_found': len(detected),
                'details': {}
            }
            for item in detected:
                t = item['type']
                stats['details'][t] = stats['details'].get(t, 0) + 1
        else:
            return {'success': False, 'error': f'未知操作: {action}'}
        
        return {
            'success': True,
            'content': result,
            'stats': stats,
            'direct_write': False,
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


# ========== Word 深度清理（参考 python-docx-cleaner）==========

def clean_docx_deep(file_path, output_path=None, options=None):
    """
    Word 文档深度清理
    
    参考 python-docx-cleaner 的思路：
    - 移除连续空段落
    - 移除孤立 run（不在段落中的 run）
    - 移除纯 bullet 字符
    - 移除非图片形状
    
    Args:
        file_path: 输入文件路径
        output_path: 输出文件路径，None 表示覆盖原文件
        options: dict, 清理选项
    
    Returns:
        dict: {success, stats, error}
    """
    opts = options or {}
    
    if not DEPS.get('docx'):
        return {'success': False, 'error': 'python-docx 未安装'}
    
    try:
        import docx
        from docx.oxml.ns import qn
        
        doc = docx.Document(file_path)
        stats = {
            'empty_paragraphs_removed': 0,
            'bullet_runs_removed': 0,
            'shapes_removed': 0,
            'orphan_runs_removed': 0
        }
        
        # 1. 移除连续空段落
        if opts.get('remove_empty_paragraphs', True):
            paragraphs_to_remove = []
            prev_empty = False
            for para in doc.paragraphs:
                is_empty = not para.text.strip()
                if is_empty and prev_empty:
                    paragraphs_to_remove.append(para)
                    stats['empty_paragraphs_removed'] += 1
                prev_empty = is_empty
            
            # 删除空段落（通过删除 XML 元素）
            for para in paragraphs_to_remove:
                p = para._element
                p.getparent().remove(p)
        
        # 2. 移除纯 bullet 字符
        if opts.get('remove_bullet_runs', True):
            bullet_chars = {'•', '-', '●', '○', '■', '□', '►', '→', '*'}
            for para in doc.paragraphs:
                runs_to_remove = []
                for run in para.runs:
                    if run.text.strip() in bullet_chars:
                        runs_to_remove.append(run)
                        stats['bullet_runs_removed'] += 1
                for run in runs_to_remove:
                    run._element.getparent().remove(run._element)
        
        # 3. 移除非图片形状
        if opts.get('remove_non_image_shapes', True):
            # 遍历所有段落中的形状
            for para in doc.paragraphs:
                for run in para.runs:
                    # 检查 run 中是否有绘图或文本框
                    drawing_elements = run._element.findall(qn('w:drawing'))
                    for drawing in drawing_elements:
                        # 检查是否是图片
                        inline = drawing.find(qn('wp:inline'))
                        if inline is None:
                            # 不是内联图片，可能是文本框或其他形状
                            drawing.getparent().remove(drawing)
                            stats['shapes_removed'] += 1
        
        # 保存
        output = output_path or file_path
        doc.save(output)
        
        return {
            'success': True,
            'stats': stats,
            'output_path': output
        }
        
    except Exception as e:
        return {'success': False, 'error': str(e)}


# ========== 举一反三：其他格式清理 ==========

def clean_txt_md(content, options=None, is_markdown=False):
    """
    清理 TXT/MD 文件 - 目标是减小体积
    
    Args:
        content: 文件内容
        options: 选项字典
        is_markdown: 是否是 MD 文件（MD 需要保留空行）
    """
    import re
    opts = options or {}
    
    if not content:
        return content
    
    result = content
    
    # 统一换行符
    result = result.replace('\r\n', '\n').replace('\r', '\n')
    
    # 移除行尾空白
    lines = [line.rstrip() for line in result.split('\n')]
    
    if is_markdown:
        # MD 文件：只移除多余的连续空行（保留段落分隔）
        # 合并 3 个以上空行为 2 个（MD 需要空行分隔段落）
        result = '\n'.join(lines)
        result = re.sub(r'\n{4,}', '\n\n\n', result)  # 3个以上空行 → 2个
    else:
        # TXT 文件：可以移除所有空行
        if opts.get('remove_empty_lines', True):
            lines = [line for line in lines if line.strip()]
        result = '\n'.join(lines)
    
    # 移除行内多余空格（两种文件都适用）
    result = re.sub(r' {2,}', ' ', result)  # 多个空格 → 1个
    
    return result


def clean_json_content(content, options=None):
    """
    清理 JSON 文件 - 目标是减小体积
    - 压缩输出（无缩进、无空格）
    """
    import json
    opts = options or {}
    
    if not content:
        return content
    
    try:
        obj = json.loads(content)
        # 紧凑输出，无缩进无空格
        return json.dumps(obj, ensure_ascii=False, separators=(',', ':'))
    except json.JSONDecodeError:
        return content  # 解析失败，返回原内容


def clean_csv_content(content, options=None):
    """
    清理 CSV 文件 - 目标是减小体积
    - 移除行尾空白
    - 移除空行
    """
    opts = options or {}
    
    if not content:
        return content
    
    lines = content.split('\n')
    cleaned = []
    
    for line in lines:
        # 移除行尾空白
        line = line.strip()
        # 保留非空行
        if line:
            cleaned.append(line)
    
    return '\n'.join(cleaned)


def clean_html_content(content, options=None):
    """
    清理 HTML 文件 - 目标是减小体积
    - 移除注释
    - 压缩空白
    - 移除冗余标签
    """
    import re
    opts = options or {}
    
    if not content:
        return content
    
    result = content
    
    # 移除 HTML 注释
    if opts.get('remove_comments', True):
        result = re.sub(r'<!--.*?-->', '', result, flags=re.DOTALL)
    
    # 压缩所有连续空白为单个空格
    if opts.get('compress_whitespace', True):
        result = re.sub(r'\s+', ' ', result)
    
    # 移除标签间多余空格
    result = re.sub(r'>\s+<', '><', result)
    
    return result.strip()


def clean_pdf_metadata(file_path, output_path=None):
    """
    清理 PDF 元数据 - 目标是减小体积
    - 移除作者、标题等元数据
    - 压缩内容流
    - 如果失败则返回原文件（不增大体积）
    """
    if not DEPS.get('pypdf'):
        return {'success': False, 'error': 'pypdf 未安装，跳过 PDF 处理'}
    
    try:
        from pypdf import PdfReader, PdfWriter
        import os
        import shutil
        
        orig_size = os.path.getsize(file_path)
        
        reader = PdfReader(file_path)
        writer = PdfWriter()
        
        # 复制所有页面
        for page in reader.pages:
            writer.add_page(page)
        
        # 清理元数据
        writer.add_metadata({
            '/Producer': '',
            '/Creator': '',
            '/Author': '',
            '/Title': '',
            '/Subject': '',
            '/Keywords': ''
        })
        
        # 保存时压缩
        output = output_path or file_path
        temp_output = output + '.tmp'
        with open(temp_output, 'wb') as f:
            writer.write(f)
        
        # 检查是否真的减小了
        new_size = os.path.getsize(temp_output)
        
        if new_size < orig_size:
            # 压缩有效，使用新文件
            os.replace(temp_output, output)
            return {'success': True, 'output_path': output, 'skipped': False, 'size_reduced': orig_size - new_size}
        else:
            # 压缩无效，复制原文件到输出路径
            os.remove(temp_output)
            if output != file_path:
                shutil.copy2(file_path, output)
            return {
                'success': True, 
                'output_path': output,
                'skipped': True,
                'reason': f'压缩无效（原:{orig_size} → 新:{new_size}），已复制原文件'
            }
        
    except Exception as e:
        return {'success': False, 'error': str(e)}


# ========== 图片压缩 ==========

def compress_image_gui(input_path, output_path=None, quality=60, max_width=None, max_height=None):
    """
    GUI 图片压缩接口

    Args:
        input_path: 输入图片路径
        output_path: 输出图片路径（None 则覆盖原文件）
        quality: JPEG 质量 (1-100)，默认 60（比旧版 85 更容易压缩）
        max_width: 最大宽度
        max_height: 最大高度

    Returns:
        dict: {'success': bool, 'output_path': str, 'original_size': int, 'new_size': int, ...}
              success=False 且 no_change=True 表示无法进一步压缩（文件已是最优）
    """
    import os, tempfile
    from PIL import Image

    try:
        orig_size = os.path.getsize(input_path)
        orig_ext = os.path.splitext(input_path)[1].lower()

        # BMP 格式不支持压缩，只能缩放
        if orig_ext == '.bmp':
            return {
                'success': False,
                'error': 'BMP 格式不支持压缩，仅支持缩放尺寸',
                'no_change': True,
                'output_path': input_path,
                'original_size': orig_size,
                'new_size': orig_size,
            }

        # 打开图片
        img = Image.open(input_path)
        orig_w, orig_h = img.width, img.height

        # 缩放尺寸（仅缩小，不放大）
        if max_width or max_height:
            max_w = max_width or orig_w
            max_h = max_height or orig_h
            max_size = (max_w, max_h)
            img.thumbnail(max_size, Image.Resampling.LANCZOS)

        # RGBA 图片：JPEG 需要转 RGB，其他格式保持 RGBA
        need_rgb = (img.mode == 'RGBA' and orig_ext in ['.jpg', '.jpeg'])
        if need_rgb:
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background

        # 输出路径
        if output_path is None:
            output_path = input_path

        ext = os.path.splitext(output_path)[1].lower()
        # 保存到临时文件，比较后再决定是否覆盖
        tmp_fd = None
        tmp_path = None
        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=ext)
            os.close(tmp_fd)
            tmp_path_orig = tmp_path

            # 按格式保存
            if ext in ['.jpg', '.jpeg']:
                img.save(tmp_path, 'JPEG', quality=quality, optimize=True)
            elif ext == '.png':
                img.save(tmp_path, 'PNG', optimize=True, compress_level=6)
            elif ext == '.webp':
                img.save(tmp_path, 'WEBP', quality=quality, method=6)
            elif ext == '.gif':
                img.save(tmp_path, 'GIF', optimize=True)
            elif ext == '.bmp':
                # BMP 不支持压缩，临时保存再比较（理论上总是变大）
                img.save(tmp_path, 'BMP')
            else:
                img.save(tmp_path, quality=quality, optimize=True)

            new_size = os.path.getsize(tmp_path)

            # 只有压缩后变小（或尺寸变化）才覆盖原文件
            if new_size < orig_size:
                # 替换原文件
                if output_path != input_path:
                    import shutil
                    shutil.move(tmp_path, output_path)
                    final_path = output_path
                else:
                    import shutil
                    shutil.move(tmp_path, output_path)
                    final_path = output_path
                return {
                    'success': True,
                    'output_path': final_path,
                    'original_size': orig_size,
                    'new_size': new_size,
                    'original_dimensions': f"{orig_w}x{orig_h}",
                    'new_dimensions': f"{img.width}x{img.height}",
                    'size_reduced': orig_size - new_size,
                    'saved_percent': round((1 - new_size/orig_size) * 100, 1) if orig_size > 0 else 0,
                    'direct_write': True,
                }
            else:
                # 压缩后没有变小，删除临时文件，返回原文件
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
                return {
                    'success': False,
                    'error': '图片已是最优状态，无法进一步压缩',
                    'no_change': True,
                    'output_path': input_path,
                    'original_size': orig_size,
                    'new_size': orig_size,
                    'original_dimensions': f"{orig_w}x{orig_h}",
                    'new_dimensions': f"{img.width}x{img.height}",
                }

        finally:
            # 清理临时文件（如果还在的话）
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    except Exception as e:
        return {'success': False, 'error': str(e)}


def get_image_info_gui(path):
    """获取图片信息"""
    from PIL import Image
    import os
    
    try:
        img = Image.open(path)
        size = os.path.getsize(path)
        
        return {
            'width': img.width,
            'height': img.height,
            'mode': img.mode,
            'format': img.format,
            'size': size,
            'size_str': format_size(size),
            'dimensions': f"{img.width} x {img.height}"
        }
    except Exception as e:
        return {'error': str(e)}


def format_size(size):
    """格式化文件大小"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


def convert_format_to_ssd(file_path, embed_images=True):
    '''将文档转换为 SSD 格式（使用 MarkItDown v2）'''
    try:
        result = convert_to_ssd_v2(file_path, embed_images=embed_images)
        return result
    except Exception as e:
        return f'转换失败: {e}'
