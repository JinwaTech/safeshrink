# -*- coding: utf-8 -*-
"""
format_to_ssd_v2.py - 使用 MarkItDown 的高保真 SSD 转换

支持格式：
- 文档：DOCX, DOC, PPTX, PPT, XLSX, PDF
- 网页：HTML, HTM
- 其他：TXT, MD, JSON, CSV, XML 等文本格式

特点：
- .doc 文件：doc2docx 转 docx，再用 python-docx 解析原始表格结构
- .docx/.pptx/.xlsx：使用 MarkItDown
- 图片自动嵌入 Base64
"""

from pathlib import Path
import logging
import re
import zipfile
import base64
import os
import tempfile

# ─── doc2docx: .doc → .docx（依赖 pywin32） ─────────────────────────────
DOC2DOCX_AVAILABLE = False
try:
    from doc2docx import convert as doc_convert
    DOC2DOCX_AVAILABLE = True
except ImportError:
    print("[WARNING] doc2docx 不可用，.doc 文件将无法转换")

# ─── MarkItDown（可选，离线降级） ──────────────────────────────────────────
try:
    from markitdown import MarkItDown
    MARKITDOWN_AVAILABLE = True
except ImportError as e:
    MarkItDown = None
    MARKITDOWN_AVAILABLE = False
    print(f"[WARNING] markitdown 不可用，禁用 SSD 转换功能: {e}")

# ─── python-docx（用于 .doc 文件的表格结构解析） ───────────────────────────
DOCX_PARSER_AVAILABLE = False
try:
    import docx
    DOCX_PARSER_AVAILABLE = True
except ImportError:
    print("[WARNING] python-docx 不可用，.doc 表格结构解析可能失败")

# ─── 支持格式定义 ───────────────────────────────────────────────────────────
SUPPORTED_FORMATS = {
    '.txt', '.ssd', '.md', '.json', '.csv', '.xml', '.yaml', '.yml',
    '.js', '.ts', '.py', '.java', '.c', '.cpp', '.h', '.css', '.sql', '.sh', '.bat',
    '.docx', '.doc', '.pptx', '.ppt', '.xlsx', '.xls',
    '.pdf', '.html', '.htm',
    '.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg', '.bmp', '.ico', '.tiff',
}

MARKDOWN_CONVERTIBLE = {
    '.docx', '.doc', '.pptx', '.ppt', '.xlsx', '.xls',
    '.pdf', '.html', '.htm',
    '.txt', '.ssd', '.md', '.json', '.csv', '.xml', '.yaml', '.yml',
}

CODE_FILES = {
    '.js', '.ts', '.py', '.java', '.c', '.cpp', '.h', '.css', '.sql', '.sh', '.bat',
}

IMAGE_EMBED_FORMATS = {'.docx', '.pptx', '.xlsx'}

IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg', '.bmp', '.ico', '.tiff'}

# ─── 全局 MarkItDown 实例 ───────────────────────────────────────────────────
_md_instance = None

def get_markitdown_instance():
    global _md_instance
    if not MARKITDOWN_AVAILABLE:
        return None
    if _md_instance is None:
        _md_instance = MarkItDown()
    return _md_instance


def is_ssd_convertible(file_path: str) -> bool:
    """检查文件是否可以转换为 SSD 格式"""
    ext = Path(file_path).suffix.lower()
    return ext in MARKDOWN_CONVERTIBLE


def get_mime_type(ext: str) -> str:
    mime_types = {
        '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.gif': 'image/gif', '.webp': 'image/webp', '.svg': 'image/svg+xml',
        '.bmp': 'image/bmp', '.ico': 'image/x-icon', '.tiff': 'image/tiff',
    }
    return mime_types.get(ext.lower(), 'image/png')


def extract_images_from_office(file_path: str) -> list:
    """从 Office 文件（DOCX/PPTX/XLSX）中提取所有图片"""
    images = []
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            for name in z.namelist():
                if 'media/' in name and not name.endswith('/'):
                    ext = os.path.splitext(name)[1].lower()
                    if ext in IMAGE_EXTENSIONS:
                        data = z.read(name)
                        images.append({
                            'name': name,
                            'data': data,
                            'mime': get_mime_type(ext),
                        })
    except Exception as e:
        print(f"    [图片提取失败] {e}")
    return images


def embed_images_in_ssd(ssd_text: str, images: list) -> str:
    """将提取的图片以 Base64 形式嵌入 SSD"""
    for i, img in enumerate(images):
        b64 = base64.b64encode(img['data']).decode('utf-8')
        data_uri = f"data:{img['mime']};base64,{b64}"
        alt = os.path.basename(img['name'])
        markdown_img = f"![]({data_uri})"
        ssd_text = ssd_text.replace(f"![]({alt})", markdown_img)
        if f"![]({alt})" not in ssd_text:
            ssd_text += f"\n{markdown_img}\n"
    return ssd_text


# ─── .doc 表格结构解析（XML 级别，处理 colspan / vMerge） ──────────────────

def _to_markdown_table(table) -> str:
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    rows_el = table._element.findall('.//w:tr', ns)
    if not rows_el:
        return ""
    tbl = table._element
    tblGrid = tbl.find('w:tblGrid', ns)
    num_cols = len(tblGrid.findall('w:gridCol', ns)) if tblGrid is not None else 0
    if num_cols == 0:
        num_cols = sum(
            int(tc.find('w:tcPr/w:gridSpan', ns).get('val', 1))
            if tc.find('w:tcPr', ns) is not None and tc.find('w:tcPr/w:gridSpan', ns) is not None
            else 1
            for tc in rows_el[0].findall('w:tc', ns)
        )
    num_rows = len(rows_el)
    grid = [[''] * num_cols for _ in range(num_rows)]
    for r_idx, tr in enumerate(rows_el):
        col_ptr = 0
        for tc in tr.findall('w:tc', ns):
            while col_ptr < num_cols and grid[r_idx][col_ptr]:
                col_ptr += 1
            if col_ptr >= num_cols:
                break
            texts = [t.text or '' for t in tc.findall('.//w:t', ns)]
            cell_text = ''.join(texts).strip()
            tcPr = tc.find('w:tcPr', ns)
            col_span = 1
            if tcPr is not None:
                gs = tcPr.find('w:gridSpan', ns)
                if gs is not None:
                    col_span = int(gs.get('val', '1'))
            end_c = min(col_ptr + col_span, num_cols)
            grid[r_idx][col_ptr] = cell_text
            col_ptr = end_c
    lines = []
    for r_idx, row in enumerate(grid):
        if not any(c.strip() for c in row):
            continue
        escaped = [c.replace('|', '\\|').replace('\n', ' ') for c in row]
        lines.append('| ' + ' | '.join(escaped) + ' |')
        if r_idx == 0:
            lines.append('| ' + ' | '.join('---' for _ in row) + ' |')
    return '\n'.join(lines)


def _is_table_only(doc) -> bool:
    """判断文档是否几乎全是表格（正文段落少于3个）"""
    content_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    return len(content_paragraphs) < 3


def convert_doc_to_ssd(file_path: str) -> str:
    """
    .doc 文件专用转换：
    1. doc2docx 转为 .docx
    2. python-docx 解析原始表格结构
    3. 输出带正确 Markdown 表格的 SSD
    """
    temp_docx_path = None
    try:
        # Step 1: .doc → .docx
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            temp_docx_path = tmp.name
        doc_convert(str(file_path), temp_docx_path)
        print(f"    [转换] .doc → .docx 成功 ({os.path.getsize(temp_docx_path)} bytes)")

        if not DOCX_PARSER_AVAILABLE:
            raise ValueError("python-docx 不可用，无法解析 .doc 表格")

        # Step 2: 读取 docx 结构
        doc = docx.Document(temp_docx_path)
        content_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        tables = doc.tables

        # 判断文档类型
        is_table_doc = _is_table_only(doc)
        print(f"    [分析] 段落数={len(content_paragraphs)}, 表格数={len(tables)}, 表格优先={is_table_doc}")

        parts = []

        # 标题段落（不在表格内的）
        if doc.paragraphs and doc.paragraphs[0].text.strip():
            title = doc.paragraphs[0].text.strip()
            parts.append(f"# {title}\n")

        # 处理每个块：段落 + 表格
        table_idx = 0
        current_paragraphs = []
        blocks = []

        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            if tag == 'p':
                # 找到对应的 python-docx 段落
                for p in doc.paragraphs:
                    if p._element == element:
                        if p.text.strip():
                            current_paragraphs.append(p.text.strip())
                        elif current_paragraphs:
                            # 空段落分隔：先输出当前段落缓冲
                            if blocks and blocks[-1].startswith('|'):
                                blocks.append('\n'.join(current_paragraphs))
                            else:
                                if current_paragraphs:
                                    blocks.append('\n'.join(current_paragraphs))
                            current_paragraphs = []
                        break
            elif tag == 'tbl' and table_idx < len(tables):
                # 输出段落缓冲
                if current_paragraphs:
                    blocks.append('\n'.join(current_paragraphs))
                    current_paragraphs = []
                # 输出表格
                md_table = _to_markdown_table(tables[table_idx])
                if md_table:
                    blocks.append(md_table)
                table_idx += 1

        if current_paragraphs:
            blocks.append('\n'.join(current_paragraphs))

        # 如果表格很多（>3个），直接用 MarkItDown 作为备选
        if len(tables) > 3 and len(content_paragraphs) > 5:
            print(f"    [混合文档] 使用 MarkItDown 作为主解析器")
            md = get_markitdown_instance()
            if md:
                result = md.convert(temp_docx_path)
                ssd_text = result.text_content if hasattr(result, 'text_content') else str(result)
                return ssd_text

        return '\n\n'.join(blocks)

    finally:
        if temp_docx_path and os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except Exception:
                pass


# ─── SSD 优化 ────────────────────────────────────────────────────────────────
def optimize_ssd(ssd_text: str) -> str:
    """去除 SSD 中的冗余内容"""
    if not ssd_text:
        return ssd_text

    # 去除连续空行
    lines = ssd_text.split('\n')
    cleaned = []
    prev_empty = False
    for line in lines:
        is_empty = not line.strip()
        if is_empty:
            if not prev_empty:
                cleaned.append('')
            prev_empty = True
        else:
            cleaned.append(line)
            prev_empty = False
    ssd_text = '\n'.join(cleaned)

    # 去除多余空格
    ssd_text = re.sub(r'[ \t]+\n', '\n', ssd_text)
    ssd_text = re.sub(r'\n{3,}', '\n\n', ssd_text)

    return ssd_text.strip()


# ─── 主转换函数 ─────────────────────────────────────────────────────────────
def convert_to_ssd_v2(file_path: str, embed_images: bool = True, optimize: bool = True) -> str:
    """
    统一转换入口：
    - .doc: 专用结构化解析（doc2docx + python-docx 表格）
    - .docx/.pptx/.xlsx: MarkItDown
    - 其他: 直接读取或 MarkItDown
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext not in SUPPORTED_FORMATS:
        raise ValueError(f"不支持的文件格式: {ext}")

    # 代码文件：直接读取
    if ext in CODE_FILES:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        lang = ext[1:]
        return f"```{lang}\n{content}\n```\n"

    temp_docx_path = None

    try:
        # ── .doc: 专用解析 ──────────────────────────────────────────────────
        if ext == '.doc':
            if not DOC2DOCX_AVAILABLE:
                raise ValueError(".doc 格式需要 doc2docx 库支持，请安装：pip install doc2docx")
            return convert_doc_to_ssd(file_path)

        # ── 图片格式：直接嵌入 ─────────────────────────────────────────────
        if ext in IMAGE_EXTENSIONS:
            with open(file_path, 'rb') as f:
                data = f.read()
            b64 = base64.b64encode(data).decode('utf-8')
            mime = get_mime_type(ext)
            return f"![{file_path.name}](data:{mime};base64,{b64})\n"

        # ── .docx/.pptx/.xlsx: 提取图片 + MarkItDown ─────────────────────
        actual_path = str(file_path)
        images = []
        if embed_images and ext in IMAGE_EMBED_FORMATS:
            images = extract_images_from_office(actual_path)
            if images:
                print(f"    [图片] 从 {ext} 文件提取了 {len(images)} 张图片")

        # MarkItDown 转换
        if not MARKITDOWN_AVAILABLE:
            raise ValueError("MarkItDown 不可用")

        md = get_markitdown_instance()
        result = md.convert(actual_path)
        ssd_text = result.text_content if hasattr(result, 'text_content') else str(result)

        if not ssd_text:
            raise ValueError("转换结果为空")

        # 嵌入图片
        if images:
            ssd_text = embed_images_in_ssd(ssd_text, images)
            print(f"    [图片] 已嵌入 {len(images)} 张图片")

        # 优化
        if optimize:
            orig_len = len(ssd_text)
            ssd_text = optimize_ssd(ssd_text)
            delta = orig_len - len(ssd_text)
            if delta > 0:
                print(f"    [优化] 去除 {delta} 字符冗余")

        return ssd_text

    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"SSD 转换失败: {e}")
