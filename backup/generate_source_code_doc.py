# -*- coding: utf-8 -*-
"""
生成软件著作权登记所需的源代码文档
前30页 + 后30页，每页50行，带页眉页脚和行号
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 源代码文件列表（按逻辑顺序排列）
SOURCE_FILES = [
    "main_window_v2.py",      # 主窗口
    "theme_manager.py",       # 主题管理
    "translations.py",        # 多语言
    "safe_shrink.py",         # 核心功能
    "safe_shrink_gui.py",     # GUI入口
    "slim_tab.py",            # 文档减肥
    "sanitize_tab.py",        # 敏感信息脱敏
    "sanitize_ssd.py",        # SSD脱敏处理
    "format_to_ssd.py",       # 格式转换
    "ssd_embed_images.py",    # 图片嵌入
    "batch_processor.py",     # 批量处理
    "batch_tab.py",           # 批量处理标签页
    "settings_tab.py",        # 设置
    "history_manager.py",     # 历史管理
    "history_tab.py",         # 历史标签页
    "file_status.py",         # 文件状态
    "result_compare_dialog.py", # 结果对比
]

LINES_PER_PAGE = 50
PAGES_REQUIRED = 30
TOTAL_LINES_FRONT = LINES_PER_PAGE * PAGES_REQUIRED  # 1500行
TOTAL_LINES_BACK = LINES_PER_PAGE * PAGES_REQUIRED   # 1500行

def read_all_source_code(base_dir):
    """读取所有源代码，合并为一个列表"""
    all_lines = []
    current_file = None
    file_boundaries = []  # 记录每个文件的起始行号
    
    for filename in SOURCE_FILES:
        filepath = Path(base_dir) / filename
        if not filepath.exists():
            print(f"警告: 文件不存在 {filepath}")
            continue
        
        file_boundaries.append((len(all_lines), filename))
        
        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            # 添加文件分隔注释
            all_lines.append(f"# {'='*60}\n")
            all_lines.append(f"# 文件: {filename}\n")
            all_lines.append(f"# {'='*60}\n\n")
            all_lines.extend(lines)
            all_lines.append("\n")  # 文件间空行
    
    return all_lines, file_boundaries

def add_page_number(section):
    """添加页码到页脚"""
    footer = section.footer
    footer.is_linked_to_previous = False
    
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加页码
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_header(section, title):
    """添加页眉"""
    header = section.header
    header.is_linked_to_previous = False
    
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run(title)
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def create_source_code_document(base_dir, output_path):
    """创建源代码文档"""
    print("读取源代码...")
    all_lines, file_boundaries = read_all_source_code(base_dir)
    total_lines = len(all_lines)
    
    print(f"总行数: {total_lines}")
    print(f"需要: 前{TOTAL_LINES_FRONT}行 + 后{TOTAL_LINES_BACK}行")
    
    # 获取前30页和后30页的代码
    front_lines = all_lines[:TOTAL_LINES_FRONT]
    
    if total_lines <= TOTAL_LINES_FRONT + TOTAL_LINES_BACK:
        # 如果代码不够，就用全部
        back_lines = all_lines[TOTAL_LINES_FRONT:]
    else:
        back_lines = all_lines[-TOTAL_LINES_BACK:]
    
    print(f"前30页: {len(front_lines)}行")
    print(f"后30页: {len(back_lines)}行")
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # 添加标题页
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SafeShrink（密小件）文档处理软件")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("源 代 码")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("V1.0.0")
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_paragraph()  # 空行
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"源程序量: 约{total_lines}行")
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 分页符
    doc.add_page_break()
    
    # 设置页眉页脚
    section = doc.sections[0]
    add_header(section, "SafeShrink（密小件）源代码 V1.0.0")
    add_page_number(section)
    
    # 添加前30页代码
    print("添加前30页...")
    add_code_section(doc, front_lines, 1, "前30页")
    
    # 分页符
    doc.add_page_break()
    
    # 添加省略说明
    omit = doc.add_paragraph()
    omit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = omit.add_run(f"... 中间省略 {total_lines - TOTAL_LINES_FRONT - TOTAL_LINES_BACK} 行 ...")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_page_break()
    
    # 添加后30页代码
    print("添加后30页...")
    start_line = total_lines - TOTAL_LINES_BACK + 1
    add_code_section(doc, back_lines, start_line, "后30页")
    
    # 保存文档
    doc.save(output_path)
    print(f"保存到: {output_path}")
    
    return total_lines

def add_code_section(doc, lines, start_line_num, section_name):
    """添加代码段"""
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(9)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    line_num = start_line_num
    
    for line in lines:
        # 创建段落
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        
        # 行号（右对齐，固定宽度）
        num_run = para.add_run(f"{line_num:5d}  ")
        num_run.font.name = 'Consolas'
        num_run.font.size = Pt(9)
        num_run.font.color.rgb = None  # 默认颜色
        
        # 代码内容
        code_run = para.add_run(line.rstrip('\n\r'))
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9)
        code_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        line_num += 1

if __name__ == "__main__":
    base_dir = Path(__file__).parent
    output_path = base_dir / "SafeShrink_源代码_V1.0.0.docx"
    
    total = create_source_code_document(base_dir, output_path)
    print(f"\n完成！总代码量: {total}行")
